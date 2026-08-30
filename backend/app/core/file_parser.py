"""Unified file parser — PDF/DOCX/MD/TXT/XLSX support."""
import os
import re
import uuid
import logging
from typing import Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".md", ".txt", ".xlsx", ".xls"}

# Chunking constants
CHUNK_MIN_SIZE = 200
CHUNK_MAX_SIZE = 2000
CHUNK_TARGET_SIZE = 800
SENTENCE_OVERLAP_COUNT = 1


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def is_supported(filename: str) -> bool:
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


def parse_file(filepath: str, filename: Optional[str] = None) -> dict:
    """Unified file parsing entry point.

    Returns: dict with keys: paper_id, title, authors, year, sections, raw_text, figures
    """
    fname = filename or os.path.basename(filepath)
    ext = get_file_extension(fname)

    if ext == ".pdf":
        return _parse_pdf(filepath, fname)
    elif ext in (".doc", ".docx"):
        return _parse_docx(filepath, fname)
    elif ext in (".md", ".txt"):
        return _parse_text(filepath, fname)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(filepath, fname)
    else:
        raise ValueError(f"Unsupported format: {ext}")


def _parse_pdf(filepath: str, filename: str) -> dict:
    """PDF parsing with pdfplumber + PyPDF2 metadata."""
    import pdfplumber
    from PyPDF2 import PdfReader

    paper_id = str(uuid.uuid4())
    metadata: dict = {"paper_id": paper_id, "filename": filename}

    try:
        reader = PdfReader(filepath)
        info = reader.metadata
        if info:
            metadata["title"] = info.title or ""
            metadata["authors"] = info.author or ""
        metadata["page_count"] = len(reader.pages)
    except Exception as e:
        logger.warning(f"Metadata extraction failed: {e}")
        metadata["page_count"] = 0

    if not metadata.get("title") or not metadata.get("authors"):
        _parse_from_filename(filename, metadata)

    sections = []
    all_figures: list[dict] = []
    try:
        with pdfplumber.open(filepath) as pdf:
            current_section: dict = {"title": "全文", "text": "", "tables": []}
            section_pattern = re.compile(
                r'^(\d+(?:\.\d+)*)\s+([A-Z\u4e00-\u9fff][^\n]{2,80})'
                r'|^[一二三四五六七八九十百]+[、．.]\s*([^\n]{2,80})'
                r'|^第[一二三四五六七八九十百]+[章节]\s*([^\n]{2,80})'
            )

            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if not text.strip():
                    continue

                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        table_md = _table_to_markdown(table)
                        current_section["tables"].append({
                            "page": i + 1,
                            "content": table_md,
                            "section_title": current_section["title"],
                        })

                figures_on_page = _extract_figures_from_page(page, i + 1, text)
                all_figures.extend(figures_on_page)

                lines = text.split("\n")
                buffer_lines: list[str] = []
                for line in lines:
                    stripped = line.strip()
                    match = section_pattern.match(stripped)
                    if match and len(stripped) < 80:
                        if buffer_lines:
                            current_section["text"] += "\n" + "\n".join(buffer_lines)
                            buffer_lines = []
                        if current_section["text"].strip():
                            sections.append(current_section.copy())
                        section_title = (
                            f"{match.group(1)} {match.group(2).strip()}" if match.group(1) and match.group(2)
                            else match.group(3).strip() if match.group(3)
                            else match.group(4).strip() if match.group(4)
                            else stripped
                        )
                        current_section = {"title": section_title, "text": "", "tables": []}
                    else:
                        buffer_lines.append(line)

                if buffer_lines:
                    current_section["text"] += "\n" + "\n".join(buffer_lines)

            if current_section["text"].strip():
                sections.append(current_section)

    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        return {**metadata, "sections": [], "raw_text": "", "figures": [], "error": str(e)}

    if not sections:
        all_text = ""
        try:
            import pdfplumber as pp
            with pp.open(filepath) as pdf:
                for page in pdf.pages:
                    all_text += (page.extract_text() or "") + "\n"
        except Exception:
            pass
        sections = [{"title": "全文", "text": all_text, "tables": []}]

    return {
        **metadata,
        "sections": sections,
        "figures": all_figures,
        "raw_text": "\n".join(s["text"] for s in sections),
    }


def _parse_docx(filepath: str, filename: str) -> dict:
    paper_id = str(uuid.uuid4())
    title = os.path.splitext(filename)[0]
    authors = ""
    year = 0

    try:
        from docx import Document

        doc = Document(filepath)
        core_props = doc.core_properties
        if core_props.title:
            title = core_props.title
        if core_props.author:
            authors = core_props.author
        if core_props.created:
            year = core_props.created.year

        full_text_parts: list[str] = []
        sections: list[dict] = []
        current_section: dict = {"title": "全文", "text": "", "tables": []}

        heading_pattern = re.compile(
            r"^(\d+(?:\.\d+)*)\s+(.+)|^([一二三四五六七八九十]+[、．.])\s*(.+)", re.UNICODE
        )

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name and "Heading" in para.style.name:
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": text, "text": "", "tables": []}
            elif heading_pattern.match(text) and len(text) < 80:
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": text, "text": "", "tables": []}
            else:
                current_section["text"] += text + "\n"
                full_text_parts.append(text)

        for table in doc.tables:
            table_md = _docx_table_to_markdown(table)
            if table_md:
                current_section["tables"].append({"content": table_md})

        if current_section["text"].strip():
            sections.append(current_section)

        raw_text = "\n".join(full_text_parts)
        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return {"paper_id": paper_id, "filename": filename, "title": title,
                "authors": authors, "year": year, "sections": [], "raw_text": "",
                "figures": [], "error": str(e)}

    return {"paper_id": paper_id, "filename": filename, "title": title,
            "authors": authors, "year": year, "sections": sections,
            "raw_text": raw_text, "figures": []}


def _parse_text(filepath: str, filename: str) -> dict:
    paper_id = str(uuid.uuid4())
    title = os.path.splitext(filename)[0]

    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

        md_heading = re.compile(r"^(#{1,6})\s+(.+)", re.MULTILINE)
        sections: list[dict] = []
        current_section: dict = {"title": "全文", "text": "", "tables": []}
        lines = raw_text.split("\n")
        buffer: list[str] = []

        for line in lines:
            match = md_heading.match(line)
            if match:
                if buffer:
                    current_section["text"] += "\n".join(buffer) + "\n"
                    buffer = []
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": match.group(2).strip(), "text": "", "tables": []}
            else:
                buffer.append(line)

        if buffer:
            current_section["text"] += "\n".join(buffer) + "\n"
        if current_section["text"].strip():
            sections.append(current_section)
        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"Text file parsing failed: {e}")
        return {"paper_id": paper_id, "filename": filename, "title": title,
                "authors": "", "year": 0, "sections": [], "raw_text": "",
                "figures": [], "error": str(e)}

    return {"paper_id": paper_id, "filename": filename, "title": title,
            "authors": "", "year": 0, "sections": sections,
            "raw_text": raw_text, "figures": []}


def _parse_xlsx(filepath: str, filename: str) -> dict:
    paper_id = str(uuid.uuid4())
    title = os.path.splitext(filename)[0]

    try:
        from openpyxl import load_workbook

        wb = load_workbook(filepath, read_only=True, data_only=True)
        sections: list[dict] = []
        full_text_parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            text_parts = [f"## Sheet: {sheet_name}"]
            header = rows[0]
            text_parts.append("| " + " | ".join(header) + " |")
            text_parts.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                padded = row + [""] * (len(header) - len(row))
                text_parts.append("| " + " | ".join(padded[: len(header)]) + " |")
            section_text = "\n".join(text_parts)
            sections.append({"title": sheet_name, "text": section_text, "tables": []})
            full_text_parts.append(section_text)
        wb.close()
        raw_text = "\n\n".join(full_text_parts)
        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"XLSX parsing failed: {e}")
        return {"paper_id": paper_id, "filename": filename, "title": title,
                "authors": "", "year": 0, "sections": [], "raw_text": "",
                "figures": [], "error": str(e)}

    return {"paper_id": paper_id, "filename": filename, "title": title,
            "authors": "", "year": 0, "sections": sections,
            "raw_text": raw_text, "figures": []}


# ─── Chunking ──────────────────────────────────────────────────────

def chunk_paper(paper_data: dict) -> list[dict]:
    """Hierarchical chunking: L0 (paper) → L1 (section) → L2 (paragraph)."""
    paper_id = paper_data["paper_id"]
    chunks: list[dict] = []

    abstract = _extract_abstract(paper_data["raw_text"])
    if abstract:
        chunks.append(_build_chunk(paper_data, "摘要", "L0", abstract))

    for section in paper_data.get("sections", []):
        text = section["text"].strip()
        if not text:
            continue
        section_title = section["title"]
        has_table = len(section.get("tables", [])) > 0

        if len(text) <= CHUNK_TARGET_SIZE:
            chunks.append(_build_chunk(paper_data, section_title, "L1", text, has_table=has_table))
        else:
            l1_chunk = _build_chunk(paper_data, section_title, "L1", text[:CHUNK_TARGET_SIZE], has_table=has_table)
            chunks.append(l1_chunk)
            sub_texts = _split_by_semantic_boundaries(text)
            for sub in sub_texts:
                chunks.append(_build_chunk(paper_data, section_title, "L2", sub, has_table=has_table,
                                           parent_chunk_id=l1_chunk["chunk_id"]))

    for section in paper_data.get("sections", []):
        section_text = section["text"].strip()
        for table in section.get("tables", []):
            context = _build_table_context(table, section_text)
            if context:
                chunks.append(_build_chunk(paper_data,
                                           f"{table.get('section_title', section['title'])} - 表格",
                                           "L2", context, has_table=True))

    logger.info(f"Paper {paper_id} chunked: {len(chunks)} chunks")
    return chunks


def _split_by_semantic_boundaries(text: str) -> list[str]:
    if not text or not text.strip():
        return []
    sentences = _split_sentences(text)
    if not sentences:
        return [text[:CHUNK_MAX_SIZE]]
    chunks = _merge_sentences_to_chunks(sentences)
    return chunks


def _split_sentences(text: str) -> list[str]:
    protected = text
    protected = re.sub(r"(\d+\.(?:\d+\.)*)(?=\s)", r"__NUMDOT__\1", protected)
    protected = re.sub(r"\b(Fig|Eq|et al|e\.g|i\.e|vs|cf|ref|al)\.", r"\1__DOT__", protected)
    parts = re.split(r"(?<=[。！？])|(?<=[.!?])(?=\s|$)", protected)
    sentences = []
    for p in parts:
        s = p.strip()
        if not s:
            continue
        s = s.replace("__NUMDOT__", "").replace("__DOT__", ".")
        sentences.append(s)
    return sentences


def _merge_sentences_to_chunks(sentences: list[str]) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for sent in sentences:
        sent_len = len(sent)
        if sent_len > CHUNK_MAX_SIZE:
            if current:
                chunks.append(" ".join(current))
                current = []
                current_len = 0
            chunks.append(sent[:CHUNK_MAX_SIZE])
            continue
        if current_len + sent_len > CHUNK_TARGET_SIZE and current:
            chunks.append(" ".join(current))
            current = [sent]
            current_len = sent_len
        else:
            current.append(sent)
            current_len += sent_len

    if current:
        text = " ".join(current)
        if len(text) < CHUNK_MIN_SIZE and chunks:
            chunks[-1] = chunks[-1] + " " + text
        else:
            chunks.append(text)
    return chunks


def _extract_abstract(text: str) -> str:
    if not text:
        return ""
    patterns = [
        r"(?:Abstract|ABSTRACT|摘要|内容摘要)[\s\n：:]*((?:.|\n){100,}?)(?=\n\s*\n|\n(?:Introduction|INTRODUCTION|1[\s.]\s|Keywords|关键词|1\s))",
        r"(?:Abstract|ABSTRACT)[\s\n：:]*(.*?)(?:\n\n|\x0c)",
        r"(?:摘要)[\s\n：:]*(.*?)(?:\n\n|关键词|Abstract)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            abstract = match.group(1).strip()
            if len(abstract) > 100:
                return abstract[:2000]
    first_page = text[:3000]
    paragraphs = re.split(r"\n\s*\n", first_page)
    for para in paragraphs:
        clean = para.strip()
        if len(clean) > 150 and not re.match(r"^\d+\.", clean):
            return clean[:1500]
    return text[:800].strip()


def _build_table_context(table: dict, section_text: str) -> str:
    table_content = table.get("content", "")
    if not table_content:
        return ""
    paragraphs = re.split(r"\n\s*\n", section_text) if section_text else []
    context_before = paragraphs[0][:200] if paragraphs else ""
    context_after = paragraphs[-1][:200] if len(paragraphs) > 1 else ""
    parts = []
    if context_before:
        parts.append(f"[上下文] {context_before}")
    parts.append(f"[表格内容]\n{table_content}")
    if context_after:
        parts.append(f"[后续内容] {context_after}")
    return "\n\n".join(parts)


def _build_chunk(paper_data: dict, section_title: str, level: str, text: str,
                 has_table: bool = False, parent_chunk_id: str = "") -> dict:
    return {
        "chunk_id": f"{paper_data['paper_id']}_{level}_{uuid.uuid4().hex[:8]}",
        "paper_id": paper_data["paper_id"],
        "project_id": paper_data.get("project_id", ""),
        "paper_title": paper_data.get("title", ""),
        "authors": paper_data.get("authors", ""),
        "year": paper_data.get("year", 0),
        "section_title": section_title,
        "section_level": level,
        "text": text,
        "has_table": has_table,
        "parent_chunk_id": parent_chunk_id,
    }


def _parse_from_filename(filename: str, metadata: dict):
    name = filename.replace(".pdf", "")
    parts = name.split(" - ")
    if len(parts) >= 3:
        metadata["authors"] = parts[0].strip()
        try:
            metadata["year"] = int(parts[1].strip())
        except ValueError:
            metadata["year"] = 0
        metadata["title"] = parts[2].strip()
    elif len(parts) == 2:
        metadata["authors"] = parts[0].strip()
        metadata["title"] = parts[1].strip()
    else:
        metadata["title"] = name


def _table_to_markdown(table: list) -> str:
    if not table or len(table) < 2:
        return ""
    header = "| " + " | ".join(str(c or "") for c in table[0]) + " |"
    separator = "| " + " | ".join("---" for _ in table[0]) + " |"
    rows = ["| " + " | ".join(str(c or "") for c in row) + " |" for row in table[1:]]
    return header + "\n" + separator + "\n" + "\n".join(rows)


def _docx_table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    data_rows = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return header + "\n" + separator + "\n" + "\n".join(data_rows)


def _extract_figures_from_page(page, page_num: int, page_text: str) -> list[dict]:
    figures: list[dict] = []
    try:
        images = page.images
    except Exception:
        images = []
    if not images:
        return []

    caption_pattern = re.compile(r"^(?:Fig\.?|Figure|图)\s*\d*[\.\:：]?\s*(.{5,})", re.IGNORECASE)
    captions: dict[int, str] = {}
    lines = page_text.split("\n")
    for j, line in enumerate(lines):
        match = caption_pattern.match(line.strip())
        if match:
            captions[j] = line.strip()

    for img in images:
        width = img.get("x1", 0) - img.get("x0", 0)
        height = img.get("bottom", 0) - img.get("top", 0)
        if width < 50 or height < 50:
            continue
        caption = ""
        img_y = img.get("top", 0)
        for line_idx, cap_text in captions.items():
            caption_y = line_idx * 12
            if abs(caption_y - img_y - height) < 100 or abs(caption_y - img_y) < 50:
                caption = cap_text
                break
        figures.append({
            "figure_id": f"fig_{uuid.uuid4().hex[:8]}",
            "page": page_num,
            "caption": caption or f"Figure on page {page_num}",
            "width": round(width, 1),
            "height": round(height, 1),
        })
    return figures
