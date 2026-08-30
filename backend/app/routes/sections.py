"""Section (draft) CRUD routes + AI generation."""

import uuid

from fastapi import APIRouter, Depends, HTTPException

from app.database import SupabaseDB
from app.deps import get_current_user, get_db, verify_project_owner
from app.models.models import Section
from app.schemas.section import (
    ExportRequest,
    SectionCreate,
    SectionGenerateRequest,
    SectionRead,
    SectionUpdate,
)

router = APIRouter(tags=["sections"])


async def _get_owned_section(db: SupabaseDB, section_id: uuid.UUID, user_id: uuid.UUID) -> dict:
    section = await db.get(Section, section_id)
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")
    await verify_project_owner(db, section["project_id"], user_id)
    return section


@router.get("/", response_model=list[SectionRead])
async def list_sections(
    project_id: str,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    await verify_project_owner(db, project_id, user_id)
    return await db.select(
        Section,
        order="order_index",
        project_id=str(uuid.UUID(project_id)),
    )


@router.post("/", response_model=SectionRead)
async def create_section(
    data: SectionCreate,
    project_id: str,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    await verify_project_owner(db, project_id, user_id)
    row = await db.insert(Section, {
        "project_id": str(uuid.UUID(project_id)),
        **data.model_dump(),
    })
    return row


@router.get("/{section_id}", response_model=SectionRead)
async def get_section(
    section_id: uuid.UUID,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    return await _get_owned_section(db, section_id, user_id)


@router.patch("/{section_id}", response_model=SectionRead)
async def update_section(
    section_id: uuid.UUID,
    data: SectionUpdate,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    await _get_owned_section(db, section_id, user_id)
    updated = await db.update(Section, str(section_id), data.model_dump(exclude_unset=True))
    return updated


@router.delete("/{section_id}", status_code=204)
async def delete_section(
    section_id: uuid.UUID,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    await _get_owned_section(db, section_id, user_id)
    await db.delete(Section, str(section_id))


@router.post("/generate")
async def generate_section(
    req: SectionGenerateRequest,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    """AI-generate section content using RAG + LLM."""
    from app.core.llm import llm_client
    from app.core.prompt import SYSTEM_PROMPT_BASE, prompt_engine
    from app.core.retriever import hybrid_search, format_chunks_with_citations

    section = await _get_owned_section(db, req.section_id, user_id)

    # RAG retrieval
    chunks = await hybrid_search(
        section["title"], top_k=8, project_id=str(section["project_id"]), intent="generate", db=db
    )
    rag_content = format_chunks_with_citations(chunks) if chunks else ""

    # Build prompt
    task_prompt = prompt_engine.build_section_prompt(
        section_name=section["title"],
        section_topic=req.instruction or section["title"],
        reference_material=rag_content,
        target_words=req.target_words,
        writing_style=req.writing_style,
    )
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT_BASE},
        {"role": "user", "content": task_prompt},
    ]

    content = await llm_client.achat(messages, temperature=0.7, max_tokens=4000)

    # Update section
    sources = [
        {"title": c.get("paper_title", ""), "citation": c.get("citation", "")}
        for c in chunks
    ]
    await db.update(Section, str(section["id"]), {
        "content": content,
        "status": "generated",
        "sources": sources,
    })

    return {
        "id": str(section["id"]),
        "content": content,
        "sources": sources,
        "word_count": len(content),
    }


@router.post("/export")
async def export_document(
    req: ExportRequest,
    user_id: uuid.UUID = Depends(get_current_user),
    db: SupabaseDB = Depends(get_db),
):
    """Export all sections as markdown or docx."""
    await verify_project_owner(db, req.project_id, user_id)
    sections = await db.select(
        Section,
        order="order_index",
        project_id=str(req.project_id),
    )

    if not sections:
        raise HTTPException(status_code=404, detail="No sections to export")

    if req.format == "docx":
        from fastapi.responses import Response
        from docx import Document
        import io

        doc = Document()
        for s in sections:
            doc.add_heading(s["title"], level=2)
            doc.add_paragraph(s.get("content") or "")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=draft.docx"},
        )

    # Default: markdown
    md_content = ""
    for s in sections:
        md_content += f"## {s['title']}\n\n{s.get('content') or ''}\n\n---\n\n"

    from fastapi.responses import Response
    return Response(
        content=md_content,
        media_type="text/markdown",
        headers={"Content-Disposition": "attachment; filename=draft.md"},
    )
